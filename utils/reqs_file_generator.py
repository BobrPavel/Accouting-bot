# --------------------------------------------------------------------------------
# Модуль создания файла с реквизитами пользователя
# --------------------------------------------------------------------------------
# Импорты
# --------------------------------------------------------------------------------

import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from common.user_reqs import Requisites

# --------------------------------------------------------------------------------
# Функция генерации1
# --------------------------------------------------------------------------------


# ---- Функция для стилизации документа ----
async def docx_file_styling(doc):
    """Применяет Times New Roman, 14 pt, чёрный цвет и межстрочный интервал 1.5 ко всем параграфам."""

    for paragraph in doc.paragraphs:
        # ---- Межстрочный интервал ----
        paragraph.paragraph_format.line_spacing = 1.5

        for run in paragraph.runs:
            # ---- Шритф ----
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # для кириллицы

            # ---- Размер и цвет ----
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)



# ---- Генерация документа ----
async def generate_requisites_docx_file(fsm_data: dict, folder_path: str) -> str:
    """
    Создание DOCX файла с реквизитами на основе данных из FSM.
    fsm_data — словарь вида:
        { "field_0": "...", "field_1": "...", ... }
    """
    # ---- Добавляем к пути название файла ----
    file_path = os.path.join(folder_path, "requisites.docx")


    # ---- Формируем документ ----
    doc = Document()

    
    doc.add_heading(fsm_data.get("field_1", ""), level=1)  # сокращённое название    
    doc.add_paragraph(fsm_data.get("field_5", "")) # Адрес (берём адрес местонахождения)
    doc.add_paragraph(f"ОГРН {fsm_data.get('field_2', '')}") # ОГРН
    doc.add_paragraph(f"ИНН/КПП {fsm_data.get('field_3', '')} / {fsm_data.get('field_4', '')}") # ИНН, КПП
    doc.add_paragraph(f"тел.: {fsm_data.get('field_7', '')}") # Телефон
    doc.add_paragraph("")  # отступ


    # ======= КАРТОЧКА ОРГАНИЗАЦИИ =======
    doc.add_heading("КАРТОЧКА ОРГАНИЗАЦИИ", level=1) 

    # Все реквизиты подряд
    for index, title in enumerate(Requisites):
        if title == "Расчетный счёт":
            break
        else:
            value = fsm_data.get(f"field_{index}", "")
            doc.add_paragraph(f"{title}: {value}")

    doc.add_paragraph("")  # отступ


    # ======= Банковские реквизиты =======
    doc.add_heading("Банковские реквизиты:", level=1) 
    doc.add_paragraph(f"р/с {fsm_data.get('field_11', '')}")
    doc.add_paragraph(f"в {fsm_data.get('field_12', '')}")
    doc.add_paragraph(f"к/с {fsm_data.get('field_13', '')}")
    doc.add_paragraph(f"БИК {fsm_data.get('field_14', '')}")

    doc.add_paragraph("")  # отступ


    doc.add_paragraph(f"{fsm_data.get('field_15', '')}")  # Подписант

    await apply_global_style(doc)
    doc.save(file_path)  # Сохраняем файл


